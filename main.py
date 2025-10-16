import re
import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docxtpl import DocxTemplate
from docx import Document


class DocumentGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор Word-документів")
        self.root.geometry("750x650")

        self.template_path = ""
        self.fields = {}

        # Заголовок
        ttk.Label(
            root,
            text="Генератор документів",
            font=("SF Pro Display", 18, "bold")
        ).pack(pady=10)

        # Кнопка вибору шаблону
        ttk.Button(
            root, text="📂 Обрати шаблон (.docx)",
            command=self.load_template
        ).pack(pady=5)

        self.template_label = ttk.Label(
            root, text="Шаблон не обрано", foreground="gray"
        )
        self.template_label.pack()

        # Рамка для полів
        self.fields_frame = ttk.Frame(root)
        self.fields_frame.pack(fill="both", expand=True, pady=10)

        # Кнопка генерації документа
        ttk.Button(
            root, text="📝 Згенерувати документ",
            command=self.generate_doc
        ).pack(pady=10)

    # ------------------------- #
    #   Зчитування змінних із шаблону
    # ------------------------- #
    def extract_variables(self, path):
        doc_tpl = DocxTemplate(path)
        docx = doc_tpl.get_docx()
        variables = set()

        def extract_from_paragraphs(paragraphs):
            text = "\n".join(p.text for p in paragraphs)
            found = re.findall(r"{{\s*([^{}]+?)\s*}}", text)
            variables.update(found)

        # 1️⃣ Основний текст
        extract_from_paragraphs(docx.paragraphs)

        # 2️⃣ Таблиці (включно з вкладеними)
        def extract_from_table(table):
            for row in table.rows:
                for cell in row.cells:
                    extract_from_paragraphs(cell.paragraphs)
                    for inner_table in cell.tables:
                        extract_from_table(inner_table)

        for table in docx.tables:
            extract_from_table(table)

        # 3️⃣ Хедери та футери
        for section in docx.sections:
            extract_from_paragraphs(section.header.paragraphs)
            extract_from_paragraphs(section.footer.paragraphs)

            # Якщо є таблиці у хедерах чи футерах
            for table in section.header.tables:
                extract_from_table(table)
            for table in section.footer.tables:
                extract_from_table(table)

        return sorted(variables)

    # ------------------------- #
    #   Завантаження шаблону
    # ------------------------- #
    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if not path:
            return
        self.template_path = path
        self.template_label.config(text=os.path.basename(path), foreground="black")

        # Зчитуємо змінні з шаблону
        vars_found = self.extract_variables(path)
        self.show_fields(vars_found)

    # ------------------------- #
    #   Відображення полів
    # ------------------------- #
    def show_fields(self, vars_list):
        for widget in self.fields_frame.winfo_children():
            widget.destroy()
        self.fields.clear()

        if not vars_list:
            ttk.Label(self.fields_frame, text="❌ У шаблоні не знайдено змінних").pack()
            return

        ttk.Label(
            self.fields_frame,
            text="Заповніть значення для змінних:",
            font=("SF Pro Display", 14)
        ).pack(pady=5)

        for var in vars_list:
            frame = ttk.Frame(self.fields_frame)
            frame.pack(fill="x", pady=3, padx=20)
            ttk.Label(frame, text=f"{var}:", width=30, anchor="w").pack(side="left")
            entry = ttk.Entry(frame, width=40)
            entry.pack(side="right", fill="x", expand=True)
            self.fields[var] = entry

    # ------------------------- #
    #   Генерація документа
    # ------------------------- #
    def generate_doc(self):
        if not self.template_path:
            messagebox.showerror("Помилка", "Спочатку оберіть шаблон!")
            return

        context = {key: entry.get() for key, entry in self.fields.items()}

        try:
            doc = DocxTemplate(self.template_path)
            doc.render(context)

            # 🗂️ Вибір папки для збереження
            save_dir = filedialog.askdirectory(title="Оберіть папку для збереження")
            if not save_dir:
                return

            base_name = (
                context.get("назва_підприємства")
                or context.get("назва_компанії")
                or "документ"
            )

            safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in base_name)
            filename = f"Рішення_ООВ_{safe_name}.docx"
            output_path = os.path.join(save_dir, filename)

            doc.save(output_path)
            messagebox.showinfo("✅ Успіх", f"Документ збережено:\n{output_path}")

        except Exception as e:
            messagebox.showerror("Помилка при створенні документа", str(e))


# ------------------------- #
#           Запуск
# ------------------------- #
if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentGeneratorApp(root)
    root.mainloop()
